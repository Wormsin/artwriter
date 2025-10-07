import os
import time
from typing import Optional
from openai import OpenAI
from docx import Document
from sqlalchemy import null
import json


def run_deep_research(
    topic: str,
    output_docx_path:str,
    prompt: str,
    model: str = "o4-mini-deep-research-2025-06-26",
    background: bool = True,
    poll_interval: int = 6,
):
    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key:
        raise ValueError("No API key found. Set OPENAI_API_KEY or pass api_key argument.")

    client = OpenAI(api_key=api_key)

    system_message = (
        "Вы опытный научный писатель. Подготовьте драфт"
        "исследовательской работы с причинно-следственными"
        "связями, логическими выводами и результирующими"
        "таблицами. Используйте академический стиль, правильное"
        "цитирование и список литературы в конце."
    )

    inputs = [
        {
            "role": "developer",
            "content": [{"type": "input_text", "text": system_message}]
        },
        {
            "role": "user",
            "content": [{"type": "input_text", "text": prompt}]
        }
    ]

    create_kwargs = dict(
        model=model,
        input=inputs,
        reasoning={"summary": "auto"},
        tools=[{"type": "web_search"}],
        include=[
            "reasoning.encrypted_content",
            "web_search_call.action.sources"
                ],
        store=False,
        max_output_tokens = 25000
    )
    if background:
        create_kwargs["background"] = True

    response = client.responses.create(**create_kwargs)

    # If background, poll until finished
    if background:
        resp_id = response.id
        status = response.status
        while status in ("queued", "in_progress"):
            time.sleep(poll_interval)
            response = client.responses.retrieve(resp_id)
            status = response.status

    error_response = response.error
    if error_response == null:
        article_text = response.output[0].content[0].text
        doc = Document()
        doc.add_heading(topic, level=1)
        doc.add_paragraph(f"Query: {prompt}")
        doc.add_paragraph("")
        for section in article_text.split("\n\n"):
            doc.add_paragraph(section.strip())
        doc.save(output_docx_path)
    
    with open("./reports/logs.txt", "a", encoding="utf-8") as f:
        f.write(str(response.created_at) + "\n")
        f.write(str(error_response) + "\n")
        f.write(str(response.reasoning.summary) + "\n")
        #f.write(str(response.usage.total_tokens) + "\n")

    
    return response


def call_openai(prompt: str, model: str = "gpt-4.1-mini", max_tokens: int = 200) -> str:
    client = OpenAI()

    response = client.chat.completions.create(
        model=model,
        messages=[{"role": "user", "content": prompt}],
        max_tokens=max_tokens
    )

    return response.choices[0].message.content
