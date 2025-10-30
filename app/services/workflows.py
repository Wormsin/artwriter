import json
import logging
from pathlib import Path
from docx import Document
from docx.shared import Pt
from typing import Dict, List, Literal
import re
import os
import shutil
import glob

# Импорты из внешних модулей
from services.gemini_api import upload_files, call_llm, upload_small_file, structured_call_llm
from services.preprompts import *
from services.schemas import *


logger = logging.getLogger(__name__)



# --- УТИЛИТЫ ---
def ensure_directory(path: Path) -> Path:
    """Проверяет существование директории и создает ее, если необходимо."""
    path.mkdir(parents=True, exist_ok=True)
    return path

def save_text(content: str, output_file: Path):
    """Сохраняет текстовый контент в файл."""
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write(content)
    logger.info(f"Текст сохранен в {output_file}")

def save_json(obj: dict | list, output_file: Path):
    """Сохраняет JSON-объект в файл."""
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(obj, f, ensure_ascii=False, indent=2)
    logger.info(f"JSON сохранен в {output_file}")

        
# --- ПОИСК ДОП ФАКТОВ ---
def expand_database(topic_path: str, llm_model_name: str) -> Path | None:
    folder_path = Path(topic_path)
    folder_path_bd = ensure_directory(folder_path / "DB")
    output_file = folder_path_bd / "db_extension.txt"

    file_paths = [str(file.resolve()) for file in folder_path_bd.iterdir() if file.is_file()]
    uploaded_files = upload_files(file_paths)
    prompt = get_stage1_prompt()
    status, response, total_tokens = call_llm(
        prompt, files=uploaded_files, 
        model_name=llm_model_name, 
        temperature=0.2, 
        web_search=False
    )
    if status != "success":
        logger.error(f"Ошибка при расширении БД: {status}")
        return None
    save_text(response, output_file)
    return status, total_tokens


# --- ПОИСК НЕОЧЕВИДНЫХ СВЯЗЕЙ ---
def connect_lenses(lens_folder: Path, output_file: Path):
    """Объединяет файлы линз в один выходной файл и удаляет папку линз."""
    try:
        txt_files = [lens_folder / f for f in os.listdir(lens_folder) if f.endswith('.txt')]
        
        if not txt_files:
            logger.warning(f"В папке '{lens_folder}' не найдено .txt файлов.")
            return

        with open(output_file, 'w', encoding='utf-8') as outfile:
            for filepath in txt_files:
                try:
                    with open(filepath, 'r', encoding='utf-8') as infile:
                        content = infile.read()
                        outfile.write(content)
                        outfile.write('\n')
                        
                except Exception as e:
                    logger.error(f"Ошибка при чтении файла {filepath}: {e}")
        logger.info(f"Файлы успешно объединены в '{output_file}'")
        shutil.rmtree(lens_folder)
    except FileNotFoundError:
        logger.error(f"Ошибка: Папка '{lens_folder}' не найдена.")
    except Exception as e:
        logger.error(f"Произошла непредвиденная ошибка: {e}")

def find_connections_main(topic_path: str, llm_model_name: str) -> Path | None:
    output_folder = ensure_directory(Path(topic_path) / "FACTS" / "ALG_MAIN" / "HYP")
    output_folder_lens = ensure_directory(output_folder / "LENS")

    folder = Path(f"{topic_path}/DB")
    file_paths = [str(file.resolve()) for file in folder.iterdir() if file.is_file()]
    uploaded_files = upload_files(file_paths)

    lens = 0
    tokens=0
    while True:
        lens += 1
        prompt = get_stage2_prompt_main(lens_num=lens)
        if not prompt:
            break
        status, response, total_tokens = call_llm(
            prompt, files=uploaded_files, 
            model_name=llm_model_name, 
            temperature=1.5,
            web_search=False
        )
        if status != "success":
            logger.error(f"Ошибка при генерации линзы {lens} (main): {status}")
            return status, tokens
        tokens+=total_tokens
        lens_file = output_folder_lens / f"lens_{lens}_main.txt"
        save_text(response, lens_file)
        
        # Обновление файлов для следующей итерации
        db_ext = Path(topic_path) / "DB" / "db_extension.txt" 
        if lens == 1:
            file_paths = [lens_file, db_ext]
            uploaded_files = upload_files(file_paths)
        else:
            uploaded_files = upload_small_file(lens_file)
    
    output_file = output_folder / "db_facts.txt"
    save_text(response, output_file)
    return status, tokens

def find_connections_blind_spots(topic_path: str, llm_model_name: str) -> Path | None:
    output_folder = ensure_directory(Path(topic_path) / "FACTS" / "ALG_BLIND" / "HYP")
    output_folder_lens = ensure_directory(output_folder / "LENS")

    folder = Path(f"{topic_path}/DB")
    file_paths = [str(file.resolve()) for file in folder.iterdir() if file.is_file()]
    db_ext = f"{topic_path}/DB/db_extension.txt" 
    file_paths.append(db_ext)
    uploaded_files = upload_files(file_paths)

    lens = 0
    tokens = 0
    while True:
        lens += 1
        prompt = get_stage2_prompt_blind_spots(lens_num=lens)
        if not prompt:
            break
        status, response, total_tokens = call_llm(
            prompt, files=uploaded_files, 
            model_name=llm_model_name, 
            temperature=1.5,
            web_search=False
        )
        if status != "success":
            logger.error(f"Ошибка при генерации линзы {lens} (blind spots): {status}")
            return status, total_tokens
        tokens += total_tokens
        lens_file = output_folder_lens / f"lens_{lens}_blind_spots.txt"
        save_text(response, lens_file)
    
    # После цикла: объединение линз
    output_file = output_folder / "db_facts.txt"
    connect_lenses(output_folder_lens, output_file)
    return status, tokens


# --- ПРОВЕРКА ГИПОТЕЗ ---
def extract_blocks(text: str, start_tag: str, end_tag: str) -> List[str]:
    pattern = re.compile(rf"\[{start_tag}\](.*?)\[{end_tag}\]", re.DOTALL)
    return [block.strip() for block in pattern.findall(text)]

def connect_check_hypothese_results(folder_lenses: Path, file_hypothesis: str, output_file: Path) -> bool:
    files_checks = [folder_lenses / f for f in os.listdir(folder_lenses) if f.endswith('.txt')]
    with open(file_hypothesis, "r", encoding="utf-8") as f:
        hypotheses = extract_blocks(f.read(), "НАЧАЛО ГИПОТЕЗЫ", "КОНЕЦ ГИПОТЕЗЫ")
    all_checks = []
    for path in files_checks:
        with open(path, "r", encoding="utf-8") as f:
            checks = extract_blocks(f.read(), "НАЧАЛО ПРОВЕРКИ", "КОНЕЦ ПРОВЕРКИ")
            all_checks.append(checks)
    num_hypotheses = len(hypotheses)
    for i, checks in enumerate(all_checks, 1):
        if len(checks) != num_hypotheses:
            logger.warning(f"В файле проверки {i} найдено {len(checks)} блоков, "
                           f"а гипотез — {num_hypotheses}.")
    merged_blocks = []
    for i in range(num_hypotheses):
        block = "[НАЧАЛО ПРОВЕРКИ ГИПОТЕЗЫ]\n"
        block += hypotheses[i] + "\n\n"
        # Добавляем все проверки для этой гипотезы
        for checks in all_checks:
            if i < len(checks):
                block += checks[i] + "\n\n"
        block += "[КОНЕЦ ПРОВЕРКИ ГИПОТЕЗЫ]\n"
        merged_blocks.append(block)

    # Сохраняем результат
    save_text("\n".join(merged_blocks), output_file)
    shutil.rmtree(folder_lenses)
    return True

def check_hypotheses(topic_path: str, llm_model_name: str, facts_type: Literal["blind_spots", "main"]) -> Path | None:
    alg_folder = "ALG_MAIN" if facts_type == "main" else "ALG_BLIND"
    output_path = ensure_directory(Path(topic_path) / "FACTS" / alg_folder / "CHECK")  
    output_path_lens = ensure_directory(output_path / "LENS")
    file_hypothesis = f"{topic_path}/FACTS/{alg_folder}/HYP/db_facts.txt"
    file_paths = [file_hypothesis]
    uploaded_files = upload_files(file_paths)
    
    tokens = 0
    lens = 0
    while True:
        lens += 1
        prompt = get_stage3_prompt(lens_num=lens)
        if not prompt:
            break
        status, response, total_tokens = call_llm(
            prompt, files=uploaded_files, 
            web_search=True, model_name=llm_model_name,
            temperature=0.2
        )
        if status != "success":
            logger.error(f"Ошибка при проверке линзы {lens}: {status}")
            return status, tokens
        tokens += total_tokens
        output_file = output_path_lens / f"db_facts_checked_{lens}.txt"
        save_text(response, output_file)
        file_paths.append(output_file)
        uploaded_files = upload_files(file_paths)
        logger.info(f"Проверенные гипотезы сохранены в {output_file}")
    
    output_file = output_path / "db_facts_checked.txt"
    connect_check_hypothese_results(
        folder_lenses=output_path_lens, 
        file_hypothesis=file_hypothesis,
        output_file=output_file
    )
    return status, tokens


# --- СОЗДАНИЕ СТРУКТУРЫ СЦЕНАРИЯ ---
def build_script_structure(topic_path: str, num_series: int, llm_model_name: str) -> Path | None:
    output_dir = ensure_directory(Path(topic_path) / "STRUCTURE")
    output_file_json = output_dir / "script_structure.json"
    output_file_txt = output_dir / "script_structure.txt"

    folder = Path(f"{topic_path}/DB")
    file_paths = [str(file.resolve()) for file in folder.iterdir() if file.is_file()]
    paths_facts = glob.glob(os.path.join(f"{topic_path}/FACTS", "ALG*/CHECK/db_facts_checked.txt"))
    file_paths.extend(paths_facts)  # Исправлено: append -> extend

    uploaded_files = upload_files(file_paths)
    prompt = get_stage4_prompt(num_series)
    status, response, total_tokens = structured_call_llm(prompt, files=uploaded_files, structure=list[ScriptStructure],
                                    max_output_tokens=65536,
                                    model_name=llm_model_name,
                                    temperature=1)
    if status != "success":
        logger.error(f"Ошибка при создании структуры сценария: {status}")
        return status, total_tokens
    try:
        scripts: list[ScriptStructure] = response.parsed
    except Exception as e:
        logger.error(f"Ошибка при парсинге ответа: {e}")
        return "error", total_tokens
    scripts_raw = [script.model_dump() for script in scripts]
    scripts_with_ids: List[ScriptStructureID] = [
    ScriptStructureID.model_validate(script) for script in scripts_raw]
    scripts_ids = [script.model_dump() for script in scripts_with_ids]
    save_json(scripts_ids, output_file_json)
    json_string = json.dumps(
        scripts_ids, 
        indent=2, 
        ensure_ascii=False
    )
    save_text(json_string, output_file_txt)
    print(f"Структура сценария сохранена в {output_file_json}")
    return status, total_tokens


# --- НАПИСАНИЕ ТЕКСТА СЦЕНАРИЯ ---
def scenario_to_docx(output_dir: str):
    output_path = Path(output_dir)
    input_file = output_path / "scenario.json"
    with open(input_file, "r", encoding="utf-8") as f:
        data = json.load(f)

    for serie in data:
        doc = Document()
        serie_title = f"{serie['serie_number']}. {serie['serie_name']}"
        doc.add_heading(serie_title, level=0)

        for chapter in serie["content"]:
            # Название главы
            doc.add_heading(f"{chapter['chapter_number']}. {chapter['chapter_name']}", level=1)

            # Описание главы (курсив)
            p_desc = doc.add_paragraph()
            run_desc = p_desc.add_run(chapter["chapter_description"])
            run_desc.italic = True
            run_desc.font.size = Pt(11)

            # Основной текст (обычный)
            p_text = doc.add_paragraph(chapter["text"])
            p_text.style = "Normal"

        file_name = f"Серия_{serie['serie_number']}.docx"
        doc.save(output_path / file_name)
        logger.info(f"Создан файл: {file_name}")

def get_chapters_per_serie_from_file(file_path: str) -> tuple[Dict[int, int], List[ScenarioStructure]]:
    file_path_obj = Path(file_path)
    if not file_path_obj.exists():
        raise FileNotFoundError(f"Файл не найден по пути: {file_path}")
    try:
        # 1. Считывание содержимого файла
        with open(file_path_obj, 'r', encoding='utf-8') as f:
            json_data_string = f.read()
        # 2. Парсинг JSON-строки в список словарей
        raw_data = json.loads(json_data_string)
        # 3. Валидация и создание объектов Pydantic
        # Если файл содержит список серий, парсим каждую
        scenario_data: list[ScenarioStructure] = [
            ScenarioStructure.model_validate(item) 
            for item in raw_data
        ]
    except (json.JSONDecodeError, ValueError) as e:
        raise ValueError(f"Ошибка валидации или некорректная структура JSON в файле: {e}")
    # 4. Расчет итогового словаря Dict[int, int]
    chapters_per_serie: Dict[int, int] = {}
    for serie in scenario_data:
        # Ключ: номер серии; Значение: длина списка content
        chapters_per_serie[serie.serie_number] = len(serie.content)
        
    return chapters_per_serie, scenario_data

def update_json_structure(topic_path: str) -> bool:
    file_json = f"{topic_path}/STRUCTURE/script_structure.json"
    file_txt = f"{topic_path}/STRUCTURE/script_structure.txt"
    try:
        with open(file_txt, 'r', encoding='utf-8') as f:
            json_string = f.read()
    except FileNotFoundError:
        logger.error(f"Файл не найден по пути: {file_txt}")
        return False
    except Exception as e:
        logger.error(f"Ошибка чтения файла {file_txt}: {e}")
        return False
    try:
        data = json.loads(json_string)
    except json.JSONDecodeError as e:
        logger.error(f"Содержимое файла '{file_txt}' не является валидным JSON. Подробности: {e}")
        return False
    try:
        with open(file_json, 'w', encoding='utf-8') as f:
            json.dump(
                data, 
                f, 
                indent=2, 
                ensure_ascii=False
            )
        logger.info(f"Содержимое из '{file_txt}' сохранено и отформатировано в '{file_json}'")
        return True
    except Exception as e:
        logger.error(f"Ошибка записи в JSON файл: {e}")
        return False

def write_script_text(topic_path: str, llm_model_name: str, temperature: float) -> Path | None:
    if not update_json_structure(topic_path=topic_path):
        return None, 0
    
    output_dir = ensure_directory(Path(topic_path) / "SCENARIO")
    output_file_json = output_dir / "scenario.json"
    
    folder_db = Path(topic_path) / "DB"
    file_paths = [str(file.resolve()) for file in folder_db.iterdir() if file.is_file()]
    paths_facts = glob.glob(os.path.join(f"{topic_path}/FACTS", "ALG*/CHECK/db_facts_checked.txt"))
    file_paths.extend(paths_facts)  # Исправлено: append -> extend
    file_paths.append(f"{topic_path}/STRUCTURE/script_structure.txt")
    uploaded_files = upload_files(file_paths)

    chapters_per_serie, scenario_data = get_chapters_per_serie_from_file(f"{topic_path}/STRUCTURE/script_structure.json")

    previous_chapter_text = ""
    tokens = 0
    for s in chapters_per_serie:
        for ch in range(1, chapters_per_serie[s] + 1):
            prompt = get_stage5_prompt(ser=s, ch=ch, previous_chapter_text=previous_chapter_text)
            status, response, total_tokens = call_llm(
                prompt, files=uploaded_files, 
                model_name=llm_model_name, 
                temperature=temperature
            )
            if status != "success":
                logger.error(f"Ошибка при написании главы {ch} серии {s}: {status}")
                return None, tokens
            tokens += total_tokens
            for serie in scenario_data:
                if serie.serie_number == s:
                    target_serie = serie
                    break
            for chapter in target_serie.content:
                if chapter.chapter_number == ch:
                    target_chapter = chapter
                    break
            target_chapter.text = response
            previous_chapter_text = response
    
    try:
        scripts = [sd.model_dump() for sd in scenario_data]
        save_json(scripts, output_file_json)
        scenario_to_docx(f"{topic_path}/SCENARIO")
        return status, tokens
    except Exception as e:
        logger.error(f"Критическая ошибка при обработке или сохранении: {e}")
        return None, tokens
